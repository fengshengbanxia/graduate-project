import random,time

from docx import Document
from docx.shared import Inches

# 接收上传文件
def upload(file):
    name = str (random.randint (10000, 99999) + time.time ()) + '.' + 'docx'
    try:  # 提前创建uploads文件夹
        with open (f'./static/assets/file/{name} ', 'wb+') as fp:
            ## 分块写入文件
            for chunk in file.chunks ():
                fp.write (chunk)
        filename = f'/static/assets/file/{name}'
        print("链接是",filename)
        return filename
    except Exception as e:
        print ("error is ", e)
        return None

# 输出word文件
def word(header,data_list):
    document = Document ()
    try:  # 提前创建uploads文件夹
        ## 头部标题
        _header = document.add_heading (header, level=6)
        ## 内部内容
        for item in data_list:
            document.add_heading (item["title"], level=2)
            # 添加一个段落
            paragraph = document.add_paragraph (item["content"])
            # 段落缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches (0.5)
        ## 保存文件
        name = str (random.randint (10000, 99999) + time.time ()) + '.' + 'docx'
        document.save(f'./static/assets/file/{name}')
        filename = f'/static/assets/file/{name}'
        return filename
    except Exception as e:
        print ("error is ", e)
        return None